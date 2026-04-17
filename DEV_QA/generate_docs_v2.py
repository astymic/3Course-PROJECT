# -*- coding: utf-8 -*-
"""
Regenerate DEV_QA output docs aligned with teammate deliverables.
Includes updated cost baseline with per-person cost calculations.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

BASE_DIR = Path(r'c:\Users\chapa\Desktop\!KHPI\PROJECT\DEV_QA')
DIAGRAMS_ARCH = BASE_DIR / "Diagrams" / "architecture"
DIAGRAMS_JIRA = BASE_DIR / "Diagrams" / "jira"
OUTPUT_DIR = BASE_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

HOURLY_RATE_PM = 450       # UAH/hr
HOURLY_RATE_BA = 400       # UAH/hr
HOURLY_RATE_DEV_QA = 500   # UAH/hr - combined Dev+QA+SA role


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r.bold = bold
    r.italic = italic
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return t

def add_image(doc, path, width=Inches(5.5)):
    if path.exists():
        doc.add_picture(str(path), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False


# ============================================================
# ARCHITECTURE REPORT (aligned with WBS + Backlog)
# ============================================================
def gen_architecture():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(11)

    # Title
    add_heading(doc, "ARCHITECTURE DIAGRAM", level=0)
    add_heading(doc, "LiLu E-Commerce MVP", level=1)
    doc.add_paragraph("")
    add_para(doc, "ВЗУТТЄВА ФАБРИКА «LiLu» — E-Commerce Platform", bold=True, size=12)
    add_para(doc, "Dev/QA: Мульков Максим  |  PM: Мурадян Руслан  |  BA: Покась Ілля")
    add_para(doc, "Команда: 3Course-PROJECT  |  Дата: 10.04.2026  |  Версія: 2.0")
    doc.add_paragraph("")

    # 1
    add_heading(doc, "1. Загальна архітектура системи (Three-Tier)", level=2)
    add_para(doc, "Система побудована за класичною трирівневою архітектурою, яка дозволяє масштабувати кожен рівень незалежно:")
    add_para(doc, "  - Клієнтський рівень (Frontend) — Next.js + Tailwind CSS. SSR для SEO каталогу.")
    add_para(doc, "  - Серверний рівень (Backend) — Node.js + Express + Prisma ORM. REST API.")
    add_para(doc, "  - Рівень даних — PostgreSQL (основна БД) + Redis (кеш) + Cloudinary (фото).")
    add_image(doc, DIAGRAMS_ARCH / "Загальна архітектура системи.png")
    doc.add_paragraph("")

    # 2
    add_heading(doc, "2. Діаграма взаємодії компонентів", level=2)
    add_para(doc, "Покупець та менеджер працюють через різні інтерфейси (каталог та адмін-панель), але спільний Backend API обробляє всі запити. Зовнішні інтеграції:")
    add_para(doc, "  - LiqPay (ПриватБанк) — онлайн-оплата карткою")
    add_para(doc, "  - Нова Пошта API — вибір відділення, розрахунок вартості, створення ТТН")
    add_para(doc, "  - Cloudinary — зберігання та оптимізація фото товарів")
    add_para(doc, "  - Tawk.to — безкоштовний popup-чат для консультацій")
    add_image(doc, DIAGRAMS_ARCH / "Діаграма взаємодії компонентів.png")
    doc.add_paragraph("")

    # 3
    add_heading(doc, "3. ER-діаграма бази даних", level=2)
    add_para(doc, "7 основних таблиць. Ключова особливість — таблиця ProductSize зберігає залишки по КОЖНОМУ розміру кожного товару (критична вимога клієнта Наталії Кобринської).")
    add_table(doc,
        ["Таблиця", "Призначення", "Ключові поля"],
        [
            ["Category", "Категорії взуття", "name, slug, season (літнє/зимове)"],
            ["Product", "Товари", "name, price, material, images[], category_id"],
            ["ProductSize", "Залишки по розмірах", "product_id, size (35-46), quantity"],
            ["User", "Користувачі", "email, phone, role (buyer/manager/admin)"],
            ["Order", "Замовлення", "status, payment_method, delivery_method, total"],
            ["OrderItem", "Позиції замовлення", "order_id, product_id, size, quantity, price"],
            ["ProductVariant", "Конструктор (Фаза 2)", "sole_type, material, lining, color"],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, "Таблиця ProductVariant НЕ реалізується в MVP, але закладена в схему для Фази 2 (конструктор взуття).", italic=True)
    add_image(doc, DIAGRAMS_ARCH / "ER-діаграма бази даних.png")
    doc.add_paragraph("")

    # 4
    add_heading(doc, "4. Потік замовлення (User Journey)", level=2)
    add_para(doc, "Повний шлях покупця від каталогу до підтвердження замовлення:")
    add_para(doc, "1. Покупець відкриває каталог -> фільтрує за розміром/категорією/ціною")
    add_para(doc, "2. Відкриває картку товару -> бачить фото, наявність розмірів, розмірну сітку")
    add_para(doc, "3. Додає товар у кошик -> оформлює замовлення (ПІБ, телефон, адреса)")
    add_para(doc, "4. Обирає оплату (LiqPay карткою або накладений платіж)")
    add_para(doc, "5. Обирає доставку (Нова Пошта відділення або кур'єр)")
    add_para(doc, "6. Отримує підтвердження з номером ТТН")
    add_image(doc, DIAGRAMS_ARCH / "Потік замовлення (User Journey).png")
    doc.add_paragraph("")

    # 5
    add_heading(doc, "5. Технологічний стек", level=2)
    add_table(doc,
        ["Компонент", "Технологія", "Обґрунтування"],
        [
            ["Frontend", "Next.js 14 (React)", "SSR для SEO, App Router, великий ecosystem"],
            ["CSS Framework", "Tailwind CSS", "Utility-first, швидка адаптивна розробка"],
            ["Backend", "Node.js + Express", "Єдина мова JS/TS для full-stack"],
            ["ORM", "Prisma", "TypeScript-типізація, автоматичні міграції"],
            ["База даних", "PostgreSQL", "ACID, надійна, підтримка JSON для варіантів"],
            ["Кешування", "Redis (Upstash)", "In-memory, швидкий доступ до залишків складу"],
            ["Фото товарів", "Cloudinary", "CDN, автоматична оптимізація, free tier 25GB"],
            ["Оплата", "LiqPay (ПриватБанк)", "Найпоширеніший платіжний шлюз в Україні"],
            ["Доставка", "Нова Пошта API v2", "Офіційний API: пошук відділень, створення ТТН"],
            ["Popup-чат", "Tawk.to", "Безкоштовний, легка інтеграція (скрипт)"],
            ["Хостинг Frontend", "Vercel (Hobby)", "Zero-config для Next.js, auto SSL"],
            ["Хостинг Backend", "Railway або Render", "Free tier достатній для MVP"],
            ["Домен", ".ua через Hostiq", "Географічна прив'язка до UA ринку"],
        ]
    )
    doc.add_paragraph("")

    # 6
    add_heading(doc, "6. Ключові архітектурні рішення", level=2)

    add_heading(doc, "6.1. Облік складу в реальному часі", level=3)
    add_para(doc, "Вимога клієнта: покупець має бачити наявність КОЖНОГО розміру без запиту через Instagram.")
    add_para(doc, "  - Окрема таблиця ProductSize (product_id + size + quantity)")
    add_para(doc, "  - При оформленні замовлення quantity зменшується атомарно (транзакція)")
    add_para(doc, "  - quantity = 0 -> розмір відображається як 'Немає в наявності'")
    add_para(doc, "  - Менеджер оновлює залишки через адмін-панель без виходу на склад")

    add_heading(doc, "6.2. Готовність до конструктора (Фаза 2)", level=3)
    add_para(doc, "Клієнт планує (Фаза 2): покупець зможе обрати підошву, матеріал верху, підклад та колір.")
    add_para(doc, "  - Таблиця ProductVariant вже закладена в ER-схему")
    add_para(doc, "  - API побудовано з урахуванням ExtensionPoint для варіантів")
    add_para(doc, "  - Міграція БД не потрібна при додаванні конструктора")

    add_heading(doc, "6.3. Безпека", level=3)
    add_para(doc, "  - Паролі: bcrypt (cost factor 12)")
    add_para(doc, "  - Авторизація: JWT (access + refresh tokens)")
    add_para(doc, "  - HTTPS: SSL від Vercel (auto) / Let's Encrypt")
    add_para(doc, "  - SQL injection: захист через Prisma ORM (parameterized queries)")
    add_para(doc, "  - XSS: React автоматично екранує output")

    add_heading(doc, "6.4. Mobile-first підхід", level=3)
    add_para(doc, "Більшість покупців LiLu приходять з Instagram -> з мобільного.")
    add_para(doc, "  - Responsive breakpoints: mobile (320px) -> tablet (768px) -> desktop (1280px)")
    add_para(doc, "  - Touch-friendly UI: великі кнопки, свайпи в галереї фото")

    # 7
    doc.add_paragraph("")
    add_heading(doc, "7. Зовнішні інтеграції (відповідно до Product Backlog)", level=2)
    add_para(doc, "Розподіл за спринтами згідно з Product Backlog BA (Покась):")
    add_table(doc,
        ["Інтеграція", "Провайдер", "User Story", "Sprint", "SP"],
        [
            ["Каталог з фото", "Cloudinary", "US-01", "Sprint 1", "5"],
            ["Фільтри (розмір, колір, категорія)", "Внутрішня логіка", "US-02", "Sprint 1", "8"],
            ["Облік складу по розмірах", "PostgreSQL + Redis", "US-05, US-06", "Sprint 1", "10"],
            ["Оплата карткою", "LiqPay API", "US-12", "Sprint 2", "8"],
            ["Накладений платіж", "Внутрішня логіка", "US-13", "Sprint 2", "3"],
            ["Доставка НП", "Нова Пошта API", "US-14, US-15", "Sprint 2", "8"],
            ["Popup-консультація", "Tawk.to", "US-16", "Sprint 2", "3"],
            ["Адмін-панель", "Next.js + Prisma", "US-17..US-19", "Sprint 1-2", "13"],
        ]
    )

    out = OUTPUT_DIR / "architecture_report.docx"
    doc.save(str(out))
    print(f"[OK] {out}")


# ============================================================
# JIRA WORKFLOW (aligned with Backlog)
# ============================================================
def gen_jira():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(11)

    add_heading(doc, "JIRA WORKFLOW", level=0)
    add_heading(doc, "LiLu E-Commerce — Scrum Board", level=1)
    doc.add_paragraph("")
    add_para(doc, "ВЗУТТЄВА ФАБРИКА «LiLu» — E-Commerce Platform", bold=True, size=12)
    add_para(doc, "Dev/QA: Мульков Максим  |  PM: Мурадян Руслан  |  BA: Покась Ілля")
    add_para(doc, "Команда: 3Course-PROJECT  |  Дата: 10.04.2026  |  Версія: 2.0")
    doc.add_paragraph("")

    # Screenshots
    screenshots = [
        ("list of epics.jpg", "1. Список Epics (List View)", "9 Epics покривають весь scope MVP згідно WBS (BA Покась). Всі мають пріоритет Medium, статус TO DO. Key проєкту: LIL."),
        ("summary.jpg", "2. Summary дошки", "37 work items (9 Epics + 28 Tasks). Всі у статусі TO DO. Типи: Task (76%), Epic (24%). Проєкт готовий до старту Sprint 1 (тиждень 3)."),
        ("timeline.jpg", "3. Timeline (Gantt View)", "Візуалізація 3-х спринтів на таймлайні: квітень-травень 2026. Відповідає Schedule Baseline (PM Мурадян): T3-T4 Sprint 1, T5-T6 Sprint 2, T7-T8 Sprint 3."),
        ("3 sprints.jpg", "4. Backlog view — 3 Спринти", "Sprint 1: Каталог + Облік складу — 12 items (20 Apr - 3 May).\nSprint 2: Кошик + Оплата + Доставка — 11 items (4 May - 17 May).\nSprint 3: Інтеграції + Тестування + Деплой — 5 items (18 May - 31 May).\nBacklog порожній — всі задачі розподілені по спринтах."),
        ("sprint detailed pop out.jpg", "5. Sprint 3 — деталізація", "Sprint 3 містить: зміна статусу замовлення (Epic: Адмін-панель), реєстрація домену .ua, налаштування хостингу (Vercel + Render), SSL/HTTPS, фінальне тестування перед запуском."),
        ("board.jpg", "6. Board View (Scrum Board)", "4 колонки: TO DO -> IN PROGRESS -> REVIEW -> DONE. WIP ліміт: 2 задачі на учасника в IN PROGRESS (згідно Quality Management Plan, BA Покась)."),
    ]

    for fn, title, desc in screenshots:
        add_heading(doc, title, level=2)
        add_para(doc, desc)
        add_image(doc, DIAGRAMS_JIRA / fn)
        doc.add_paragraph("")

    # Workflow
    add_heading(doc, "7. Workflow (потік статусів задач)", level=2)
    add_para(doc, "Відповідає Quality Management Plan (QA process, BA Покась):")
    add_table(doc,
        ["З", "До", "Хто переводить", "Умова (Definition of Done)"],
        [
            ["TO DO", "IN PROGRESS", "Виконавець", "Виконавець взяв задачу в роботу"],
            ["IN PROGRESS", "REVIEW", "Виконавець", "Self-review + unit-тести пройдено"],
            ["REVIEW", "DONE", "BA (Покась) або PM", "Acceptance Criteria виконані, Code Review пройдено"],
            ["REVIEW", "IN PROGRESS", "BA або PM", "Знайдено зауваження — повернення на доопрацювання"],
        ]
    )
    doc.add_paragraph("")

    # Epics aligned with WBS
    add_heading(doc, "8. Epics відповідно до WBS", level=2)
    add_para(doc, "Структура Epics побудована на основі WBS (BA Покась, v1.0):")
    add_table(doc,
        ["Epic", "WBS ID", "User Stories", "Story Points", "Sprint"],
        [
            ["Каталог товарів", "3.1-3.4", "US-01..US-04 (4 stories)", "24 SP", "Sprint 1"],
            ["Облік складу", "3.5-3.6", "US-05..US-06 (2 stories)", "10 SP", "Sprint 1"],
            ["Авторизація", "3.7-3.8", "US-07..US-08 (2 stories)", "8 SP", "Sprint 1"],
            ["Кошик та замовлення", "4.1-4.3", "US-09..US-11 (3 stories)", "15 SP", "Sprint 2"],
            ["Оплата", "4.4-4.5", "US-12..US-13 (2 stories)", "11 SP", "Sprint 2"],
            ["Доставка", "4.6-4.7", "US-14..US-15 (2 stories)", "8 SP", "Sprint 2"],
            ["Popup-консультація", "5.1", "US-16 (1 story)", "3 SP", "Sprint 2"],
            ["Адмін-панель", "5.2-5.4", "US-17..US-19 (3 stories)", "13 SP", "Sprint 1-2"],
            ["Інфраструктура", "6.1-6.4", "US-20..US-23 (4 stories)", "10 SP", "Sprint 3"],
        ]
    )
    add_para(doc, "Всього: 23 User Stories, 102 Story Points, 3 спринти по 2 тижні.", bold=True)
    doc.add_paragraph("")

    # Sprint distribution
    add_heading(doc, "9. Розподіл Story Points по спринтах", level=2)
    add_table(doc,
        ["Sprint", "Назва", "Дати (Schedule Baseline)", "Stories", "Story Points"],
        [
            ["Sprint 1", "Каталог + Облік складу + Авторизація", "T3-T4 (20.04 - 03.05)", "12", "~55 SP"],
            ["Sprint 2", "Кошик + Оплата + Доставка + Popup", "T5-T6 (04.05 - 17.05)", "8", "~37 SP"],
            ["Sprint 3", "Інфра + Тестування + Деплой", "T7-T8 (18.05 - 31.05)", "3+", "~10 SP + bugs"],
        ]
    )
    doc.add_paragraph("")

    # Labels
    add_heading(doc, "10. Мітки (Labels)", level=2)
    add_para(doc, "Відповідно до Product Backlog (BA Покась) — MoSCoW пріоритизація:")
    add_table(doc,
        ["Label", "Кількість", "Опис"],
        [
            ["must-have", "17 stories", "Обов'язково для MVP. Без цього продукт не запуститься"],
            ["should-have", "4 stories", "Важливо, але MVP може запуститися без цього"],
            ["could-have", "2 stories", "Бажано, додається якщо залишається час у спринті"],
            ["phase-2", "—", "Заплановано на Фазу 2 (конструктор взуття)"],
        ]
    )
    doc.add_paragraph("")

    # Rules
    add_heading(doc, "11. Правила роботи з дошкою", level=2)
    rules = [
        "Кожна Story має Acceptance Criteria (визначено BA) — без них задача не приймається в Done.",
        "WIP ліміт: макс. 2 задачі в IN PROGRESS на одного учасника.",
        "Daily Standup: щодня до 10:00 у Telegram групі — 3 пункти (вчора / сьогодні / блокери).",
        "Баги (Critical, High) пріоритетніші за нові Story — блокують Demo та реліз.",
        "Sprint Review (Demo клієнту) — наприкінці T4, T6, T8 (згідно Comm Plan, PM Мурадян).",
        "Sprint Retrospective — після кожного Demo, обговорення що покращити.",
        "Definition of Done (DoD) — згідно Quality Management Plan (BA Покась).",
    ]
    for i, r in enumerate(rules, 1):
        add_para(doc, f"{i}. {r}")

    out = OUTPUT_DIR / "jira_workflow_report.docx"
    doc.save(str(out))
    print(f"[OK] {out}")


# ============================================================
# COST BASELINE (with per-person cost calculation)
# ============================================================
def gen_cost():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(11)

    add_heading(doc, "COST BASELINE", level=0)
    add_heading(doc, "LiLu E-Commerce MVP", level=1)
    doc.add_paragraph("")
    add_para(doc, "ВЗУТТЄВА ФАБРИКА «LiLu» — E-Commerce Platform", bold=True, size=12)
    add_para(doc, "Dev/QA: Мульков Максим  |  PM: Мурадян Руслан  |  BA: Покась Ілля")
    add_para(doc, "Команда: 3Course-PROJECT  |  Дата: 10.04.2026  |  Версія: 2.0")
    doc.add_paragraph("")
    add_para(doc, 'Наталія Кобринська: "Прорахуйте мені вартість цього сайту і скажіть, як довго нам чекати до запуску."', italic=True)
    doc.add_paragraph("")

    # 1
    add_heading(doc, "1. Вартість інфраструктури (щомісячні витрати)", level=2)
    add_table(doc,
        ["Компонент", "Провайдер", "Вартість (грн/міс)", "Примітка"],
        [
            ["Домен .ua", "Hostiq / Ukraine.com.ua", "~21 (250/рік)", "Реєстрація на 1 рік"],
            ["SSL-сертифікат", "Let's Encrypt (Vercel)", "0", "Автоматичне поновлення"],
            ["Хостинг Frontend", "Vercel (Hobby план)", "0", "Безкоштовний для Next.js"],
            ["Хостинг Backend", "Railway / Render", "0 - 200", "Free tier достатній для MVP"],
            ["БД PostgreSQL", "Render / Supabase", "0 - 150", "Free: 500MB (достатньо)"],
            ["Файлове сховище", "Cloudinary", "0", "Free: 25GB bandwidth"],
            ["Кешування Redis", "Upstash", "0", "Free: 10K запитів/день"],
            ["Popup-чат", "Tawk.to", "0", "Повністю безкоштовний"],
            ["РАЗОМ інфраструктура", "", "21 - 371", ""],
        ]
    )
    doc.add_paragraph("")

    # 2
    add_heading(doc, "2. Вартість інструментів", level=2)
    add_table(doc,
        ["Інструмент", "Призначення", "Вартість"],
        [
            ["Jira (Free)", "Трекінг задач, Scrum Board (до 10 осіб)", "0 грн"],
            ["GitHub", "Репозиторій, документація", "0 грн"],
            ["Figma (Free)", "Прототипування UI", "0 грн"],
            ["VS Code", "Редактор коду", "0 грн"],
            ["Node.js, Next.js, Prisma", "Стек розробки (open source)", "0 грн"],
            ["РАЗОМ інструменти", "", "0 грн"],
        ]
    )
    doc.add_paragraph("")

    # 3
    add_heading(doc, "3. Транзакційні витрати (зовнішні сервіси)", level=2)
    add_table(doc,
        ["Сервіс", "Модель оплати", "Вартість"],
        [
            ["LiqPay (ПриватБанк)", "Комісія з кожної транзакції", "2.75% від суми платежу"],
            ["Нова Пошта API", "Безкоштовний API", "0 (клієнт оплачує доставку)"],
            ["Кур'єрська доставка", "Визначає фабрика", "Фіксована ціна по місту"],
        ]
    )
    doc.add_paragraph("")

    # 4 - MAIN: per-person cost calculation
    add_heading(doc, "4. Розрахункова вартість роботи команди", level=2)
    add_para(doc, "Оцінка годин базується на WBS v1.0 (BA Покась) — 175 загальних годин на проєкт.", bold=True)
    doc.add_paragraph("")

    # 4.1 PM
    add_heading(doc, "4.1. PM — Мурадян Руслан", level=3)
    pm_tasks = [
        ["Initiation", "Charter, Stakeholders Matrix", "5 год"],
        ["Planning", "WBS review, Gantt/Schedule, Risk Register, Comm Plan", "12 год"],
        ["Sprint 1", "Sprint Planning, Daily Standup, Sprint Review/Demo", "8 год"],
        ["Sprint 2", "Sprint Planning, Daily Standup, Sprint Review/Demo", "8 год"],
        ["Sprint 3", "Sprint Review, контроль деплою, ескалації", "8 год"],
        ["Closure", "Lessons Learned, Closure Report, передача клієнту", "5 год"],
        ["Комунікація", "Telegram/Viber з клієнтом (Наталія), координація", "10 год"],
        ["РАЗОМ PM", "", "56 год"],
    ]
    add_table(doc,
        ["Фаза", "Задачі (згідно WBS / Comm Plan)", "Годин"],
        pm_tasks
    )
    add_para(doc, f"Розрахункова вартість PM: 56 год x {HOURLY_RATE_PM} грн/год = {56 * HOURLY_RATE_PM:,} грн", bold=True)
    doc.add_paragraph("")

    # 4.2 BA
    add_heading(doc, "4.2. BA — Покась Ілля", level=3)
    ba_tasks = [
        ["Initiation", "Business Case (аналіз транскрипції клієнта)", "4 год"],
        ["Planning", "WBS, Product Backlog (23 US), Quality Plan, Lifecycle", "18 год"],
        ["Sprint 1", "Acceptance testing каталогу та складу, контент", "8 год"],
        ["Sprint 2", "Acceptance testing кошика, оплати, доставки", "8 год"],
        ["Sprint 3", "Acceptance testing деплою, фінальна перевірка", "5 год"],
        ["Closure", "Підготовка UAT з клієнтом, документація", "4 год"],
        ["Контент", "Наповнення каталогу (фото, описи, ціни від клієнта)", "8 год"],
        ["РАЗОМ BA", "", "55 год"],
    ]
    add_table(doc,
        ["Фаза", "Задачі (згідно WBS / Quality Plan)", "Годин"],
        ba_tasks
    )
    add_para(doc, f"Розрахункова вартість BA: 55 год x {HOURLY_RATE_BA} грн/год = {55 * HOURLY_RATE_BA:,} грн", bold=True)
    doc.add_paragraph("")

    # 4.3 Dev/QA
    add_heading(doc, "4.3. Dev/QA — Мульков Максим", level=3)
    dev_tasks = [
        ["Planning", "Архітектура, Cost Baseline, Jira Workflow", "8 год"],
        ["Sprint 1", "Каталог (US-01..04), Облік складу (US-05..06), Авторизація (US-07..08), Адмін CRUD (US-17..18)", "28 год"],
        ["Sprint 2", "Кошик (US-09..11), Оплата LiqPay (US-12..13), НП API (US-14..15), Popup (US-16), Адмін замовлення (US-19)", "24 год"],
        ["Sprint 3", "Деплой (домен, хостинг, SSL), фінальне тестування, bug-fixes", "14 год"],
        ["QA (паралельно)", "Unit-тести, self-review, Sprint QA, регресія", "12 год"],
        ["DevOps", "CI/CD, Vercel/Render налаштування, моніторинг", "4 год"],
        ["РАЗОМ Dev/QA", "", "90 год"],
    ]
    add_table(doc,
        ["Фаза", "Задачі (згідно Backlog + Architecture)", "Годин"],
        dev_tasks
    )
    add_para(doc, f"Розрахункова вартість Dev/QA: 90 год x {HOURLY_RATE_DEV_QA} грн/год = {90 * HOURLY_RATE_DEV_QA:,} грн", bold=True)
    doc.add_paragraph("")

    # 4.4 Summary
    add_heading(doc, "4.4. Зведена таблиця вартості людських ресурсів", level=3)
    pm_cost = 56 * HOURLY_RATE_PM
    ba_cost = 55 * HOURLY_RATE_BA
    dev_cost = 90 * HOURLY_RATE_DEV_QA
    total_cost = pm_cost + ba_cost + dev_cost
    total_hours = 56 + 55 + 90

    add_table(doc,
        ["Роль", "Учасник", "Годин", "Ставка (грн/год)", "Вартість (грн)"],
        [
            ["PM", "Мурадян Руслан", "56", f"{HOURLY_RATE_PM}", f"{pm_cost:,}"],
            ["BA", "Покась Ілля", "55", f"{HOURLY_RATE_BA}", f"{ba_cost:,}"],
            ["Dev/QA/SA", "Мульков Максим", "90", f"{HOURLY_RATE_DEV_QA}", f"{dev_cost:,}"],
            ["РАЗОМ", "3 учасники", f"{total_hours}", "—", f"{total_cost:,}"],
        ]
    )
    doc.add_paragraph("")
    add_para(doc, f"Загальна вартість людських ресурсів: {total_cost:,} грн", bold=True, size=12)
    add_para(doc, "* У межах навчального проєкту людські ресурси не тарифікуються. Ставки вказані для ринкової оцінки проєкту на рівні Junior/Middle спеціалістів в Україні.", italic=True)
    doc.add_paragraph("")

    # 5 - Total
    add_heading(doc, "5. Загальна вартість запуску MVP", level=2)
    add_table(doc,
        ["Категорія", "Одноразово (грн)", "Щомісяця (грн)"],
        [
            ["Домен .ua (1 рік)", "~250", "—"],
            ["Хостинг (free tiers)", "0", "0 - 371"],
            ["Інструменти (Jira, GitHub...)", "0", "0"],
            ["Зовнішні сервіси (LiqPay)", "0", "2.75% від продажів"],
            ["Людські ресурси (ринкова оцінка)", f"{total_cost:,}", "—"],
            ["РАЗОМ (з ринковою оцінкою HR)", f"~{total_cost + 250:,}", "0 - 371"],
            ["РАЗОМ (навчальний проєкт)", "~250", "0 - 371"],
        ]
    )
    doc.add_paragraph("")

    add_heading(doc, "Висновок для клієнта:", level=3)
    add_para(doc, "1. Інфраструктурні витрати на запуск MVP: ~250 грн (домен) + 0-371 грн/міс (хостинг на free tiers).", bold=True)
    add_para(doc, f"2. Ринкова вартість розробки (якби наймали команду): ~{total_cost:,} грн ({total_hours} годин x 3 спеціалісти).", bold=True)
    add_para(doc, "3. Час до запуску: 2 місяці (8 тижнів: 1 Initiation + 1 Planning + 6 розробка).", bold=True)
    add_para(doc, "4. При масштабуванні (конструктор, більше трафіку) вартість хостингу зросте до ~$10-20/міс.", italic=True)
    doc.add_paragraph("")

    # 6 - Risks
    add_heading(doc, "6. Ризики бюджету (згідно Risk Register, PM Мурадян)", level=2)
    add_table(doc,
        ["Ризик", "Ймовірність", "Вплив", "Стратегія"],
        [
            ["Free tier хостингу стане недостатнім", "Середня", "Низький", "Перехід на платний план ($5-7/міс)"],
            ["LiqPay змінить умови комісії", "Низька", "Середній", "Альтернатива: Fondy, WayForPay"],
            ["Cloudinary free вичерпається", "Середня", "Низький", "Оптимізація фото перед завантаженням"],
            ["Перевитрата годин Dev (scope creep)", "Середня (R-01)", "Високий", "Контроль через Sprint Review + Charter"],
        ]
    )

    out = OUTPUT_DIR / "cost_baseline_report.docx"
    doc.save(str(out))
    print(f"[OK] {out}")


# ============================================================
if __name__ == "__main__":
    print("[START] Regenerating documents aligned with teammate deliverables...\n")
    gen_architecture()
    gen_jira()
    gen_cost()
    print(f"\n[DONE] All files saved to: {OUTPUT_DIR}")
