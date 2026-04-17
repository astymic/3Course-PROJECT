# -*- coding: utf-8 -*-
"""
Read teammate files and regenerate DEV_QA output docs
"""
import os
import sys
import io

# Force UTF-8 output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from pathlib import Path

folder = Path(r'c:\Users\chapa\Desktop\!KHPI\PROJECT\New Files\Files from my teametes')

for f in sorted(folder.glob('*.docx')):
    doc = Document(str(f))
    print(f'=== {f.name} ===')
    for para in doc.paragraphs[:50]:
        t = para.text.strip()
        if t:
            print(f'  {t[:150]}')
    print()
    for ti, table in enumerate(doc.tables[:5]):
        print(f'  [Table {ti+1}: {len(table.rows)}r x {len(table.columns)}c]')
        for ri, row in enumerate(table.rows[:5]):
            cells = [c.text.strip().replace('\n',' ')[:60] for c in row.cells]
            print(f'    R{ri}: {cells}')
        if len(table.rows) > 5:
            print(f'    ... +{len(table.rows)-5} rows')
    print('---\n')
