# -*- coding: utf-8 -*-
"""
Sprint 1 Report for Stakeholders - LiLu E-Commerce
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime

OUTPUT_DIR = Path(r'c:\Users\chapa\Desktop\!KHPI\PROJECT\DEV_QA\output')
OUTPUT_DIR.mkdir(exist_ok=True)


def set_cell_bg(cell, hex_color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color='1A1A2E'):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(
            int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        )
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(
            int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        )
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    return p


def add_table(doc, headers, rows, header_color='1A1A2E', header_text_color='FFFFFF'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(
            int(header_text_color[0:2], 16),
            int(header_text_color[2:4], 16),
            int(header_text_color[4:6], 16)
        )

    # Data rows
    for ri, row in enumerate(rows):
        bg = 'F8F8F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    return t


def gen_sprint1_report():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ─── TITLE PAGE ───────────────────────────────────────────
    doc.add_paragraph("")
    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ВЗУТТЄВА ФАБРИКА «LiLu»")
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("E-Commerce Platform")
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor(0x5A, 0x5A, 0x7A)

    doc.add_paragraph("")

    rep_title = doc.add_paragraph()
    rep_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = rep_title.add_run("ЗВІТ ПО ВИКОНАНІЙ РОБОТІ")
    rt.font.size = Pt(20)
    rt.bold = True
    rt.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sprint_label = doc.add_paragraph()
    sprint_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sl = sprint_label.add_run("Sprint 1 — Тижні 3–4")
    sl.font.size = Pt(14)
    sl.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    sl.bold = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Meta table (centered)
    meta = doc.add_table(rows=6, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Команда:", "3Course-PROJECT"),
        ("PM:", "Мурадян Руслан"),
        ("BA:", "Покась Ілля"),
        ("Dev/QA:", "Мульков Максим"),
        ("Дата звіту:", "17.04.2026"),
        ("Версія:", "1.0"),
    ]
    for ri, (label, value) in enumerate(meta_data):
        c0 = meta.rows[ri].cells[0]
        c1 = meta.rows[ri].cells[1]
        set_cell_bg(c0, 'F0F0F5')
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True; r0.font.size = Pt(10); r0.font.name = 'Calibri'
        r1 = c1.paragraphs[0].add_run(value)
        r1.font.size = Pt(10); r1.font.name = 'Calibri'

    doc.add_page_break()

    # ─── 1. EXECUTIVE SUMMARY ────────────────────────────────
    add_heading(doc, "1. Короткий підсумок (Executive Summary)", level=1)
    add_para(doc,
        "За перші два тижні активної розробки (Sprint 1) команда успішно реалізувала "
        "ядро інтернет-магазину LiLu. Усі заплановані функції виконані в повному обсязі "
        "без переносу на наступний спринт.",
        size=11
    )
    doc.add_paragraph("")

    add_table(doc,
        ["Показник", "Результат"],
        [
            ["Заплановано User Stories", "12 (Sprint 1)"],
            ["Виконано User Stories", "12 / 12 (100%)"],
            ["Story Points виконано", "~55 SP"],
            ["Критичних помилок (Critical/High bugs)", "0"],
            ["Статус спринту", "ЗАВЕРШЕНО"],
            ["Готовність до Demo", "ТАК"],
        ],
        header_color='2E7D32',
    )

    doc.add_paragraph("")

    # ─── 2. ЩО ЗРОБЛЕНО ─────────────────────────────────────
    add_heading(doc, "2. Що було зроблено (Deliverables)", level=1)

    # 2.1 Catalog
    add_heading(doc, "2.1. Каталог товарів з фільтрами", level=2)
    add_para(doc,
        "Реалізовано головну сторінку магазину — каталог взуття з повним набором фільтрів "
        "для зручного пошуку товарів покупцем."
    )
    doc.add_paragraph("")

    add_table(doc,
        ["User Story", "Acceptance Criteria", "Статус"],
        [
            ["US-01: Переглянути каталог із фото",
             "Сітка товарів з фото, назвою, ціною та категорією",
             "DONE"],
            ["US-02: Фільтр за розміром",
             "Позначаються лише доступні розміри; товари без розміру зникають",
             "DONE"],
            ["US-03: Фільтр за категорією",
             "Літнє / Зимове / Демісезонне — моментальне оновлення",
             "DONE"],
            ["US-04: Фільтр за ціною",
             "Поля «від» і «до» фільтрують без перезавантаження сторінки",
             "DONE"],
            ["US-05: Фільтр за кольором",
             "Кнопки-чіпси кольорів; активний колір підсвічується",
             "DONE"],
        ],
        header_color='1565C0',
    )

    doc.add_paragraph("")
    add_para(doc, "Технічна реалізація:", bold=True)
    add_bullet(doc, "Next.js 16 (App Router) — серверний рендеринг для SEO каталогу")
    add_bullet(doc, "API route GET /api/products?category=&size=&color=&minPrice=&maxPrice=")
    add_bullet(doc, "Prisma ORM + SQLite (для локальної розробки); готово до PostgreSQL")
    add_bullet(doc, "Tailwind CSS — адаптивна сітка: 1-2-3 колонки залежно від розміру екрана")

    doc.add_paragraph("")

    # 2.2 Stock
    add_heading(doc, "2.2. Система обліку залишків у реальному часі", level=2)
    add_para(doc,
        "Ключова вимога клієнта (Наталія Кобринська): «Покупець має бачити наявність "
        "кожного розміру без дзвінка чи повідомлення в Instagram.» — реалізована повністю."
    )
    doc.add_paragraph("")

    add_table(doc,
        ["User Story", "Acceptance Criteria", "Статус"],
        [
            ["US-06: Наявність розмірів на картці",
             "Розмір без залишків — сірий із закресленням; з залишком — активний",
             "DONE"],
            ["US-07: Залишки в адмін-панелі",
             "Таблиця: розмір | кількість з кольоровим кодуванням",
             "DONE"],
            ["US-08: Оновлення залишків через адмін",
             "Менеджер змінює кількість → видно покупцю миттєво",
             "DONE"],
        ],
        header_color='1565C0',
    )

    doc.add_paragraph("")
    add_para(doc, "Як це працює:", bold=True)
    add_bullet(doc, "Таблиця ProductSize: product_id + size (35–46) + quantity")
    add_bullet(doc, "quantity = 0 → розмір «Немає в наявності» (перекреслено для покупця)")
    add_bullet(doc, "quantity 1–2 → «Мало» (помаранчевий індикатор для менеджера)")
    add_bullet(doc, "quantity > 2 → «Є в наявності» (зелений індикатор)")
    add_bullet(doc, "При оновленні через адмін-панель зміни видно миттєво — без кешу")

    doc.add_paragraph("")

    # 2.3 Admin
    add_heading(doc, "2.3. Адмін-панель управління товарами", level=2)
    add_para(doc,
        "Менеджер фабрики LiLu може повністю керувати каталогом без участі розробника: "
        "додавати нові моделі, редагувати залишки по розмірах, ховати неактуальні товари."
    )
    doc.add_paragraph("")

    add_table(doc,
        ["User Story", "Acceptance Criteria", "Статус"],
        [
            ["US-09: Додати новий товар",
             "Форма: назва, опис, ціна, категорія, матеріал, колір, fото URL",
             "DONE"],
            ["US-10: Редагувати товар та залишки",
             "Поля кількості для кожного розміру 35–46 в одній формі",
             "DONE"],
            ["US-11: Видалити / приховати товар",
             "Кнопка «Видалити» + чекбокс «Активний» для приховування",
             "DONE"],
            ["US-12: Дашборд залишків",
             "4 верхніх картки: всього товарів, активних, одиниць, мало залишку",
             "DONE"],
        ],
        header_color='1565C0',
    )

    doc.add_paragraph("")
    add_para(doc, "Функції адмін-панелі:", bold=True)
    add_bullet(doc, "Захищений маршрут: /admin")
    add_bullet(doc, "Статистика складу: загальна кількість одиниць на складі")
    add_bullet(doc, "Попередження: виділення товарів з кількістю <= 2 шт. (помаранчевим)")
    add_bullet(doc, "CRUD операції через REST API: POST/PUT/DELETE /api/products")

    doc.add_paragraph("")

    # ─── 3. TECHNICAL STACK ──────────────────────────────────
    add_heading(doc, "3. Технічна імплементація", level=1)

    add_table(doc,
        ["Компонент", "Технологія", "Призначення"],
        [
            ["Frontend", "Next.js 16 + React 19", "SSR, App Router, UI"],
            ["Стилізація", "Tailwind CSS 4", "Адаптивний дизайн"],
            ["Backend", "Next.js API Routes", "REST API (products, categories)"],
            ["ORM", "Prisma 5", "Типізований доступ до БД"],
            ["БД (dev)", "SQLite", "Локальна розробка без налаштувань"],
            ["БД (prod)", "PostgreSQL (готово)", "Render / Supabase для деплою"],
            ["Мова", "TypeScript", "Типобезпека, автодоповнення"],
        ],
        header_color='1A1A2E',
    )

    doc.add_paragraph("")

    add_heading(doc, "API Endpoints (реалізовані)", level=2)
    add_table(doc,
        ["Метод", "Endpoint", "Опис"],
        [
            ["GET",    "/api/products", "Список товарів з фільтрами (category, size, color, minPrice, maxPrice)"],
            ["POST",   "/api/products", "Створити новий товар із залишками по розмірах"],
            ["GET",    "/api/products/:id", "Один товар з усіма даними"],
            ["PUT",    "/api/products/:id", "Оновити товар + залишки по розмірах"],
            ["DELETE", "/api/products/:id", "Видалити товар (cascade: видаляє і ProductSize)"],
            ["GET",    "/api/categories", "Список категорій взуття"],
        ],
        header_color='37474F',
    )

    doc.add_paragraph("")

    # ─── 4. SCHEMA ───────────────────────────────────────────
    add_heading(doc, "4. Структура бази даних (Sprint 1)", level=1)

    add_table(doc,
        ["Таблиця", "Поля", "Призначення"],
        [
            ["Category",    "id, name, slug, season",                           "Категорії взуття (літнє/зимове/демісезонне)"],
            ["Product",     "id, name, description, price, material, color,\nimageUrl, isActive, categoryId, createdAt", "Товари каталогу"],
            ["ProductSize", "id, productId, size (35-46), quantity",            "Залишки по кожному розміру кожного товару"],
        ],
        header_color='4A148C',
    )

    doc.add_paragraph("")
    add_para(doc,
        "Примітка: таблиці Order, OrderItem, User та ProductVariant будуть додані у Sprint 2 "
        "(кошик, оплата, доставка) та Фазі 2 (конструктор взуття).",
        italic=True, size=10
    )

    doc.add_paragraph("")

    # ─── 5. QA ───────────────────────────────────────────────
    add_heading(doc, "5. Тестування (QA)", level=1)

    add_table(doc,
        ["Тип тестування", "Що перевірялось", "Результат"],
        [
            ["Функціональне (ручне)",
             "Всі фільтри каталогу, CRUD адмін-панелі, відображення залишків",
             "Пройдено — 0 критичних помилок"],
            ["API тестування",
             "Всі 6 endpoints: GET/POST/PUT/DELETE з різними параметрами",
             "HTTP 200/201 — коректні відповіді"],
            ["Acceptance Testing (BA)",
             "Перевірка відповідності Acceptance Criteria кожної User Story",
             "12/12 US прийнято"],
            ["Адаптивність UI",
             "Мобільний, планшет, десктоп",
             "Коректне відображення на всіх"],
            ["Крайні випадки (Edge Cases)",
             "Порожній каталог, фільтр без результатів, товар з 0 залишків",
             "Коректна поведінка — skeleton + «Не знайдено»"],
        ],
        header_color='1B5E20',
    )

    doc.add_paragraph("")

    # ─── 6. DEMO ─────────────────────────────────────────────
    add_heading(doc, "6. Інструкція для Demo", level=1)
    add_para(doc,
        "Для перегляду результатів Sprint 1 потрібно запустити локальну версію сайту:"
    )
    doc.add_paragraph("")

    add_table(doc,
        ["Крок", "Команда / Дія"],
        [
            ["1. Перейти в папку", "cd sprint1"],
            ["2. Запустити сервер", "npm run dev"],
            ["3. Відкрити каталог", "http://localhost:3000"],
            ["4. Відкрити адмін", "http://localhost:3000/admin"],
            ["5. Спробувати фільтри", "Натисніть на категорію, колір або розмір у лівій панелі"],
            ["6. Оновити залишок", "Адмін → «Редагувати» → змінити кількість розміру → «Зберегти зміни»"],
            ["7. Перевірити каталог", "На головній сторінці побачите оновлену наявність розмірів"],
        ],
        header_color='E65100',
    )

    doc.add_paragraph("")

    # ─── 7. NEXT SPRINT ──────────────────────────────────────
    add_heading(doc, "7. Наступні кроки (Sprint 2, Тижні 5–6)", level=1)
    add_para(doc,
        "За поточним Schedule Baseline (PM Мурадян Руслан), Sprint 2 стартує з 04.05.2026 "
        "і охоплює:"
    )
    doc.add_paragraph("")

    add_table(doc,
        ["Функція", "Epic / User Stories", "Пріоритет"],
        [
            ["Кошик покупця",          "US-09..US-11 | Кошик та замовлення", "Must Have"],
            ["Оформлення замовлення",  "US-14 | Форма з ПІБ, телефон, адреса", "Must Have"],
            ["Оплата LiqPay",          "US-12 | Інтеграція LiqPay (ПриватБанк)", "Must Have"],
            ["Накладений платіж",      "US-13 | Замовлення зі статусом «очікує оплату»", "Must Have"],
            ["Доставка Нова Пошта",    "US-14, US-15 | API НП — вибір відділення", "Must Have"],
            ["Popup-чат (Tawk.to)",    "US-16 | Консультація в реальному часі", "Must Have"],
            ["Список замовлень (адмін)", "US-19 | Фільтр і статус замовлень", "Must Have"],
        ],
        header_color='1A1A2E',
    )

    doc.add_paragraph("")

    # ─── 8. RISKS ────────────────────────────────────────────
    add_heading(doc, "8. Ризики та рекомендації клієнту", level=1)

    add_table(doc,
        ["Ризик", "Вплив", "Рекомендація"],
        [
            ["Клієнт не надає фото товарів до Sprint 2",
             "Каталог без реальних фото для Demo",
             "Надати мінімум 3-5 фото на товар до 01.05.2026 (BA Покась отримає деталі)"],
            ["LiqPay тестовий режим",
             "Оплата лише з тестовою карткою",
             "Надати доступ до LiqPay merchant кабінету для отримання API ключів"],
            ["Нова Пошта API квота",
             "Обмеження на безкоштовні запити",
             "Поточні ліміти достатні для MVP — моніторинг після запуску"],
        ],
        header_color='B71C1C',
    )

    doc.add_paragraph("")

    # ─── 9. METRICS ──────────────────────────────────────────
    add_heading(doc, "9. Метрики Sprint 1", level=1)

    add_table(doc,
        ["Метрика", "Значення"],
        [
            ["Тривалість спринту",     "2 тижні (Т3–Т4)"],
            ["User Stories заплановано", "12"],
            ["User Stories виконано",   "12 (100%)"],
            ["Story Points",            "~55 SP"],
            ["API endpoints",           "6"],
            ["Компонентів UI",          "6 (Page, FilterPanel, ProductCard, AdminPage, ProductTable, ProductForm)"],
            ["Таблиць БД",              "3 (Category, Product, ProductSize)"],
            ["Тестових записів у БД",   "3 категорії, 6 товарів, 36 розмірних записів"],
            ["Critical/High bugs",      "0"],
            ["Часу витрачено (Dev/QA)", "~28 годин"],
        ],
        header_color='1A1A2E',
    )

    doc.add_paragraph("")

    # ─── FOOTER / SIGNATURE ──────────────────────────────────
    add_heading(doc, "10. Затвердження звіту", level=1)

    add_table(doc,
        ["Роль", "Учасник", "Підпис", "Дата"],
        [
            ["Dev/QA (Автор звіту)", "Мульков Максим",  "_________________", "17.04.2026"],
            ["PM (Перевірив)",       "Мурадян Руслан",  "_________________", "17.04.2026"],
            ["BA (Acceptance Test)", "Покась Ілля",     "_________________", "17.04.2026"],
        ],
        header_color='37474F',
    )

    doc.add_paragraph("")
    add_para(doc,
        "Документ підготовлено командою проєкту 3Course-PROJECT для стейкхолдера "
        "— власниці взуттєвої фабрики «LiLu» Наталії Кобринської.",
        italic=True, size=10
    )

    # Save
    out = OUTPUT_DIR / "Sprint1_Report_Stakeholder.docx"
    doc.save(str(out))
    print(f"[OK] {out}")


if __name__ == '__main__':
    gen_sprint1_report()
