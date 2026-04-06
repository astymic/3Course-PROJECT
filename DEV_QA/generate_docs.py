# -*- coding: utf-8 -*-
"""
Генерація MD (з base64-картинками) та DOCX файлів
для артефактів Dev/QA — LiLu E-Commerce
"""

import base64
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = Path(__file__).parent
DIAGRAMS_ARCH = BASE_DIR / "Diagrams" / "architecture"
DIAGRAMS_JIRA = BASE_DIR / "Diagrams" / "jira"
OUTPUT_DIR = BASE_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)


def img_to_base64(path):
    """Конвертує зображення в base64 data URI."""
    ext = path.suffix.lower().replace('.', '')
    if ext == 'jpg':
        ext = 'jpeg'
    with open(path, 'rb') as f:
        data = base64.b64encode(f.read()).decode('utf-8')
    return f"data:image/{ext};base64,{data}"


def add_styled_heading(doc, text, level=1):
    """Додає стилізований заголовок."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return heading


def add_styled_paragraph(doc, text, bold=False, italic=False, size=11):
    """Додає стилізований параграф."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return para


def add_table(doc, headers, rows):
    """Додає таблицю з заголовками."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Дані
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table


def add_image_safe(doc, path, width=Inches(5.5)):
    """Додає зображення якщо файл існує."""
    if path.exists():
        doc.add_picture(str(path), width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False


# ============================================================
# 1. ARCHITECTURE — MD з base64 + DOCX
# ============================================================
def generate_architecture_md():
    """Генерує architecture_report.md з base64-картинками."""
    parts = []
    parts.append("# Архітектурна схема системи — LiLu E-Commerce MVP\n")
    parts.append("**Автор:** Мульков Максим (Dev/QA, частково SA)  ")
    parts.append("**Дата:** 06.04.2026  ")
    parts.append("**Версія:** 2.0\n")
    parts.append("---\n")

    # Секція 1
    parts.append("## 1. Загальна архітектура системи (Three-Tier)\n")
    parts.append("Система побудована за класичною **трирівневою архітектурою**:\n")
    parts.append("- **Клієнтський рівень** (Frontend) — Next.js + Tailwind CSS")
    parts.append("- **Серверний рівень** (Backend) — Node.js + Express + Prisma ORM")
    parts.append("- **Рівень даних** — PostgreSQL + Redis + Cloudinary\n")

    img1 = DIAGRAMS_ARCH / "Загальна архітектура системи.png"
    if img1.exists():
        parts.append(f"![Загальна архітектура системи]({img_to_base64(img1)})\n")

    # Секція 2
    parts.append("---\n")
    parts.append("## 2. Діаграма взаємодії компонентів\n")
    parts.append("Показує зв'язки між Frontend, Backend, зовнішніми API (LiqPay, Нова Пошта, Cloudinary) та базою даних.\n")

    img2 = DIAGRAMS_ARCH / "Діаграма взаємодії компонентів.png"
    if img2.exists():
        parts.append(f"![Діаграма взаємодії компонентів]({img_to_base64(img2)})\n")

    # Секція 3
    parts.append("---\n")
    parts.append("## 3. ER-діаграма бази даних\n")
    parts.append("Основні сутності: **Category**, **Product**, **ProductSize** (облік складу по розмірах), **User**, **Order**, **OrderItem**.\n")
    parts.append("> ⚠️ Таблиця `ProductVariant` закладена для Фази 2 (конструктор взуття), але НЕ реалізується в MVP.\n")

    img3 = DIAGRAMS_ARCH / "ER-діаграма бази даних.png"
    if img3.exists():
        parts.append(f"![ER-діаграма бази даних]({img_to_base64(img3)})\n")

    # Секція 4
    parts.append("---\n")
    parts.append("## 4. Потік замовлення (User Journey)\n")
    parts.append("Повний шлях покупця: каталог → картка товару → кошик → оформлення → оплата (LiqPay або накладений платіж) → доставка (Нова Пошта / кур'єр).\n")

    img4 = DIAGRAMS_ARCH / "Потік замовлення (User Journey).png"
    if img4.exists():
        parts.append(f"![Потік замовлення]({img_to_base64(img4)})\n")

    # Секція 5 — Стек
    parts.append("---\n")
    parts.append("## 5. Обраний технологічний стек\n")
    parts.append("| Компонент | Технологія | Обґрунтування |")
    parts.append("|---|---|---|")
    parts.append("| Frontend | Next.js (React) | SSR для SEO каталогу |")
    parts.append("| CSS | Tailwind CSS | Швидка розробка адаптивного UI |")
    parts.append("| Backend | Node.js + Express | Єдина мова JS/TS |")
    parts.append("| ORM | Prisma | Типізація, міграції |")
    parts.append("| База даних | PostgreSQL | Надійна, безкоштовна |")
    parts.append("| Кешування | Redis (Upstash) | Швидкий доступ до залишків |")
    parts.append("| Фото | Cloudinary | Зберігання та оптимізація |")
    parts.append("| Оплата | LiqPay | Найпоширеніший в Україні |")
    parts.append("| Доставка | Нова Пошта API | Офіційний API |")
    parts.append("| Popup-чат | Tawk.to | Безкоштовний |")
    parts.append("| Хостинг | Vercel + Render | Free tiers для MVP |")
    parts.append("| Домен | .ua (Hostiq) | Географічна прив'язка |\n")

    # Секція 6 — Ключові рішення
    parts.append("---\n")
    parts.append("## 6. Ключові архітектурні рішення\n")
    parts.append("### 6.1. Облік складу в реальному часі")
    parts.append("- Наявність у таблиці `ProductSize` (запис для кожного розміру).")
    parts.append("- При замовленні `quantity` зменшується автоматично.")
    parts.append("- `quantity = 0` → розмір «Немає в наявності».\n")
    parts.append("### 6.2. Підготовка до конструктора (Фаза 2)")
    parts.append("- Таблиця `ProductVariant` закладена в ER-схему.")
    parts.append("- Комбінації: підошва × матеріал × підклад × колір.\n")
    parts.append("### 6.3. Безпека")
    parts.append("- bcrypt для паролів, JWT для авторизації.")
    parts.append("- HTTPS (SSL від Vercel / Let's Encrypt).\n")
    parts.append("### 6.4. Mobile-first")
    parts.append("- Адаптивний дизайн: мобільний → планшет → десктоп.\n")

    # Секція 7 — Інтеграції
    parts.append("---\n")
    parts.append("## 7. Зовнішні інтеграції\n")
    parts.append("| Інтеграція | Провайдер | Що робить | Sprint |")
    parts.append("|---|---|---|:---:|")
    parts.append("| Оплата карткою | LiqPay | Онлайн-платежі | 2 |")
    parts.append("| Накладений платіж | Власна логіка | Оплата при отриманні | 2 |")
    parts.append("| Доставка НП | Нова Пошта API | Вибір відділення, ТТН | 2 |")
    parts.append("| Кур'єр | Власна логіка | Фіксована ціна | 2 |")
    parts.append("| Popup-чат | Tawk.to | Зв'язок з менеджером | 2 |")
    parts.append("| Фото товарів | Cloudinary | Оптимізація зображень | 1 |\n")

    content = "\n".join(parts)
    out_path = OUTPUT_DIR / "architecture_report.md"
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"[OK] {out_path}")


def generate_architecture_docx():
    """Генерує architecture_report.docx з картинками."""
    doc = Document()

    # Стиль шрифту за замовчуванням
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Титулка
    add_styled_heading(doc, "Архітектурна схема системи", level=0)
    add_styled_heading(doc, "LiLu E-Commerce MVP", level=1)
    add_styled_paragraph(doc, "Автор: Мульков Максим (Dev/QA, частково SA)", bold=True)
    add_styled_paragraph(doc, "Дата: 06.04.2026 | Версія: 2.0")
    doc.add_paragraph("")

    # 1. Загальна архітектура
    add_styled_heading(doc, "1. Загальна архітектура системи (Three-Tier)", level=2)
    add_styled_paragraph(doc, "Система побудована за класичною трирівневою архітектурою:")
    add_styled_paragraph(doc, "• Клієнтський рівень (Frontend) — Next.js + Tailwind CSS")
    add_styled_paragraph(doc, "• Серверний рівень (Backend) — Node.js + Express + Prisma ORM")
    add_styled_paragraph(doc, "• Рівень даних — PostgreSQL + Redis + Cloudinary")
    add_image_safe(doc, DIAGRAMS_ARCH / "Загальна архітектура системи.png")
    doc.add_paragraph("")

    # 2. Взаємодія компонентів
    add_styled_heading(doc, "2. Діаграма взаємодії компонентів", level=2)
    add_styled_paragraph(doc, "Показує зв'язки між Frontend, Backend, зовнішніми API (LiqPay, Нова Пошта, Cloudinary) та базою даних.")
    add_image_safe(doc, DIAGRAMS_ARCH / "Діаграма взаємодії компонентів.png")
    doc.add_paragraph("")

    # 3. ER-діаграма
    add_styled_heading(doc, "3. ER-діаграма бази даних", level=2)
    add_styled_paragraph(doc, "Основні сутності: Category, Product, ProductSize (облік складу по розмірах), User, Order, OrderItem.")
    add_styled_paragraph(doc, "⚠️ Таблиця ProductVariant закладена для Фази 2 (конструктор взуття), але НЕ реалізується в MVP.", italic=True)
    add_image_safe(doc, DIAGRAMS_ARCH / "ER-діаграма бази даних.png")
    doc.add_paragraph("")

    # 4. User Journey
    add_styled_heading(doc, "4. Потік замовлення (User Journey)", level=2)
    add_styled_paragraph(doc, "Повний шлях покупця: каталог → картка товару → кошик → оформлення → оплата → доставка.")
    add_image_safe(doc, DIAGRAMS_ARCH / "Потік замовлення (User Journey).png")
    doc.add_paragraph("")

    # 5. Стек
    add_styled_heading(doc, "5. Обраний технологічний стек", level=2)
    add_table(doc,
        ["Компонент", "Технологія", "Обґрунтування"],
        [
            ["Frontend", "Next.js (React)", "SSR для SEO каталогу"],
            ["CSS", "Tailwind CSS", "Швидка розробка адаптивного UI"],
            ["Backend", "Node.js + Express", "Єдина мова JS/TS"],
            ["ORM", "Prisma", "Типізація, міграції"],
            ["База даних", "PostgreSQL", "Надійна, безкоштовна"],
            ["Кешування", "Redis (Upstash)", "Швидкий доступ до залишків"],
            ["Фото", "Cloudinary", "Зберігання та оптимізація"],
            ["Оплата", "LiqPay (ПриватБанк)", "Найпоширеніший в Україні"],
            ["Доставка", "Нова Пошта API", "Офіційний API"],
            ["Popup-чат", "Tawk.to", "Безкоштовний"],
            ["Хостинг", "Vercel + Render", "Free tiers для MVP"],
            ["Домен", ".ua (Hostiq)", "Географічна прив'язка"],
        ]
    )
    doc.add_paragraph("")

    # 6. Ключові рішення
    add_styled_heading(doc, "6. Ключові архітектурні рішення", level=2)

    add_styled_heading(doc, "6.1. Облік складу в реальному часі", level=3)
    add_styled_paragraph(doc, "• Наявність у таблиці ProductSize (запис для кожного розміру кожного товару).")
    add_styled_paragraph(doc, "• При замовленні quantity зменшується автоматично.")
    add_styled_paragraph(doc, "• quantity = 0 → розмір показується як «Немає в наявності».")

    add_styled_heading(doc, "6.2. Підготовка до конструктора (Фаза 2)", level=3)
    add_styled_paragraph(doc, "• Таблиця ProductVariant закладена в ER-схему, але не реалізується в MVP.")
    add_styled_paragraph(doc, "• Комбінації: підошва × матеріал × підклад × колір.")

    add_styled_heading(doc, "6.3. Безпека", level=3)
    add_styled_paragraph(doc, "• Паролі — bcrypt хешування. JWT-токени для авторизації.")
    add_styled_paragraph(doc, "• HTTPS обов'язково (SSL від Vercel / Let's Encrypt).")

    add_styled_heading(doc, "6.4. Mobile-first", level=3)
    add_styled_paragraph(doc, "• Більшість покупців приходять з Instagram → з телефону.")
    add_styled_paragraph(doc, "• Адаптивний дизайн: мобільний → планшет → десктоп.")

    # 7. Інтеграції
    doc.add_paragraph("")
    add_styled_heading(doc, "7. Зовнішні інтеграції", level=2)
    add_table(doc,
        ["Інтеграція", "Провайдер", "Що робить", "Sprint"],
        [
            ["Оплата карткою", "LiqPay", "Онлайн-платежі", "2"],
            ["Накладений платіж", "Власна логіка", "Оплата при отриманні", "2"],
            ["Доставка НП", "Нова Пошта API", "Вибір відділення, ТТН", "2"],
            ["Кур'єр", "Власна логіка", "Фіксована ціна", "2"],
            ["Popup-чат", "Tawk.to", "Зв'язок з менеджером", "2"],
            ["Фото товарів", "Cloudinary", "Оптимізація зображень", "1"],
        ]
    )

    out_path = OUTPUT_DIR / "architecture_report.docx"
    doc.save(str(out_path))
    print(f"[OK] {out_path}")


# ============================================================
# 2. JIRA WORKFLOW — MD з base64 + DOCX
# ============================================================
def generate_jira_md():
    """Генерує jira_workflow_report.md з base64-скріншотами."""
    parts = []
    parts.append("# Jira Workflow — LiLu E-Commerce\n")
    parts.append("**Автор:** Мульков Максим (Dev/QA)  ")
    parts.append("**Дата:** 06.04.2026  ")
    parts.append("**Версія:** 2.0\n")
    parts.append("---\n")

    # Дошка
    parts.append("## 1. Налаштована Jira дошка\n")
    parts.append("Проєкт **Shoe Factory E-Commerce** створено на Jira (Scrum Board). Key проєкту: `LIL`.\n")

    for name, title, desc in [
        ("list of epics.jpg", "Список Epics (List View)", "9 Epics покривають весь scope MVP: від каталогу до інфраструктури."),
        ("summary.jpg", "Summary дошки", "37 робочих елементів (9 Epics + 28 Tasks). Всі у статусі TO DO. Готові до старту Sprint 1."),
        ("timeline.jpg", "Timeline (Gantt View)", "Візуалізація 3-х спринтів на таймлайні: квітень–травень 2026."),
        ("3 sprints.jpg", "Backlog — 3 Спринти", "Sprint 1: Каталог + Облік складу (12 items). Sprint 2: Кошик + Оплата + Доставка (11 items). Sprint 3: Інтеграції + Тест + Деплой (5 items)."),
        ("sprint detailed pop out.jpg", "Sprint 3 — деталі", "Деталізація Sprint 3: зміна статусу замовлення, домен, хостинг, SSL, фінальне тестування."),
        ("board.jpg", "Board View", "Дошка з колонками To Do → In Progress → Review → Done."),
    ]:
        img_path = DIAGRAMS_JIRA / name
        parts.append(f"### {title}\n")
        parts.append(f"{desc}\n")
        if img_path.exists():
            parts.append(f"![{title}]({img_to_base64(img_path)})\n")
        else:
            parts.append(f"*⚠️ Файл `{name}` не знайдено*\n")

    # Workflow опис
    parts.append("---\n")
    parts.append("## 2. Workflow (потік статусів)\n")
    parts.append("```")
    parts.append("TO DO → IN PROGRESS → REVIEW → DONE")
    parts.append("                         ↓")
    parts.append("                   IN PROGRESS (якщо зауваження)")
    parts.append("```\n")
    parts.append("| З → До | Хто переводить | Умова |")
    parts.append("|---|---|---|")
    parts.append("| To Do → In Progress | Виконавець | Взяв задачу |")
    parts.append("| In Progress → Review | Виконавець | Результат готовий |")
    parts.append("| Review → Done | PM або QA | Перевірено |")
    parts.append("| Review → In Progress | PM або QA | Зауваження |\n")

    # Epics та Stories
    parts.append("---\n")
    parts.append("## 3. Структура Epics та Stories (28 шт.)\n")

    epics_data = [
        ("📦 Каталог товарів", ["Список товарів з фото", "Фільтр за розміром", "Фільтр за категорією", "Фільтр за ціною", "Фільтр за кольором"]),
        ("👟 Картка товару", ["Фото з усіх боків", "Наявність розмірів", "Розмірна сітка", "Опис та матеріал"]),
        ("📊 Облік складу", ["Залишки по розмірах", "Оновлення в адмін-панелі", "Авто-зменшення при замовленні"]),
        ("🛒 Кошик та замовлення", ["Додати в кошик", "Оформити замовлення", "Вибір доставки", "Вибір оплати"]),
        ("💳 Оплата", ["LiqPay інтеграція", "Накладений платіж"]),
        ("🚚 Доставка", ["Нова Пошта API", "Кур'єрська доставка"]),
        ("💬 Popup-консультація", ["Tawk.to віджет"]),
        ("🔧 Адмін-панель", ["CRUD товарів", "Список замовлень", "Статус замовлення"]),
        ("⚙️ Інфраструктура", ["Домен .ua", "Хостинг", "SSL", "Фінальне тестування"]),
    ]

    for epic_name, stories in epics_data:
        parts.append(f"### {epic_name}")
        for s in stories:
            parts.append(f"- {s}")
        parts.append("")

    # Спринти
    parts.append("---\n")
    parts.append("## 4. Розподіл по спринтах\n")
    parts.append("| Sprint | Назва | Дати | Items |")
    parts.append("|---|---|---|:---:|")
    parts.append("| 1 | Каталог + Облік складу | 20.04 → 03.05 | 12 |")
    parts.append("| 2 | Кошик + Оплата + Доставка | 04.05 → 17.05 | 11 |")
    parts.append("| 3 | Інтеграції + Тест + Деплой | 18.05 → 31.05 | 5 |\n")

    content = "\n".join(parts)
    out_path = OUTPUT_DIR / "jira_workflow_report.md"
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"[OK] {out_path}")


def generate_jira_docx():
    """Генерує jira_workflow_report.docx зі скріншотами."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_styled_heading(doc, "Jira Workflow", level=0)
    add_styled_heading(doc, "LiLu E-Commerce — Scrum Board", level=1)
    add_styled_paragraph(doc, "Автор: Мульков Максим (Dev/QA)", bold=True)
    add_styled_paragraph(doc, "Дата: 06.04.2026 | Версія: 2.0")
    doc.add_paragraph("")

    # Скріншоти
    screenshots = [
        ("list of epics.jpg", "1. Список Epics (List View)", "9 Epics покривають весь scope MVP: від каталогу до інфраструктури. Всі мають пріоритет Medium та статус TO DO."),
        ("summary.jpg", "2. Summary дошки", "37 робочих елементів: 9 Epics + 28 Tasks. Всі у статусі TO DO — проєкт готовий до старту Sprint 1. Типи: Task (76%), Epic (24%)."),
        ("timeline.jpg", "3. Timeline (Gantt View)", "Візуалізація 3-х спринтів: квітень → травень → початок червня 2026. Кожен Epic має свій бар на таймлайні відповідно до спринту."),
        ("3 sprints.jpg", "4. Backlog — 3 Спринти", "Sprint 1: Каталог + Облік складу — 12 елементів (20 Apr – 3 May).\nSprint 2: Кошик + Оплата + Доставка — 11 елементів (4 May – 17 May).\nSprint 3: Інтеграції + Тест + Деплой — 5 елементів (18 May – 31 May).\nBacklog порожній — всі задачі розподілені."),
        ("sprint detailed pop out.jpg", "5. Sprint 3 — деталі", "Деталізація Sprint 3: зміна статусу замовлення (Адмін-панель), реєстрація домену, налаштування хостингу (Vercel + Render), SSL (HTTPS), фінальне тестування перед запуском."),
        ("board.jpg", "6. Board View", "Дошка з 4 колонками: TO DO → IN PROGRESS → REVIEW → DONE. WIP ліміт: 2 задачі на людину в IN PROGRESS."),
    ]

    for filename, title, description in screenshots:
        add_styled_heading(doc, title, level=2)
        add_styled_paragraph(doc, description)
        img_path = DIAGRAMS_JIRA / filename
        add_image_safe(doc, img_path)
        doc.add_paragraph("")

    # Workflow
    add_styled_heading(doc, "7. Workflow (потік статусів задач)", level=2)
    add_styled_paragraph(doc, "Кожна задача проходить 4 статуси:")
    add_table(doc,
        ["З", "До", "Хто переводить", "Умова"],
        [
            ["TO DO", "IN PROGRESS", "Виконавець", "Взяв задачу в роботу"],
            ["IN PROGRESS", "REVIEW", "Виконавець", "Результат готовий для перевірки"],
            ["REVIEW", "DONE", "PM або QA", "Перевірено, прийнято"],
            ["REVIEW", "IN PROGRESS", "PM або QA", "Зауваження — доопрацювати"],
        ]
    )
    doc.add_paragraph("")

    # Типи тікетів
    add_styled_heading(doc, "8. Типи тікетів", level=2)
    add_table(doc,
        ["Тип", "Опис", "Приклад"],
        [
            ["Epic", "Великий блок функціоналу", "Каталог товарів"],
            ["Story", "User Story — функція з точки зору користувача", "Як покупець, я хочу фільтрувати за розміром"],
            ["Task", "Технічна або документальна задача", "Налаштувати PostgreSQL"],
            ["Bug", "Дефект знайдений при тестуванні", "Popup не відкривається"],
        ]
    )
    doc.add_paragraph("")

    # Спринти
    add_styled_heading(doc, "9. Розподіл по спринтах", level=2)
    add_table(doc,
        ["Sprint", "Назва", "Дати", "Елементів"],
        [
            ["Sprint 1", "Каталог + Облік складу", "20.04 → 03.05", "12"],
            ["Sprint 2", "Кошик + Оплата + Доставка", "04.05 → 17.05", "11"],
            ["Sprint 3", "Інтеграції + Тест + Деплой", "18.05 → 31.05", "5"],
        ]
    )
    doc.add_paragraph("")

    # Мітки
    add_styled_heading(doc, "10. Мітки (Labels)", level=2)
    add_table(doc,
        ["Label", "Використання"],
        [
            ["must-have", "Критичний функціонал MVP"],
            ["should-have", "Важливо, але можна відкласти"],
            ["could-have", "Бонусний функціонал"],
            ["phase-2", "Фаза 2 (конструктор)"],
        ]
    )

    # Правила
    doc.add_paragraph("")
    add_styled_heading(doc, "11. Правила роботи з дошкою", level=2)
    rules = [
        "Кожна Story має Acceptance Criteria — без них задача не переходить в DONE.",
        "WIP ліміт: макс. 2 задачі в IN PROGRESS на одного учасника.",
        "Daily Update: кожен оновлює статус тікетів щодня.",
        "Баги пріоритетніші за нові Story.",
        "Sprint Retrospective: наприкінці кожного спринту обговорення що покращити.",
    ]
    for i, rule in enumerate(rules, 1):
        add_styled_paragraph(doc, f"{i}. {rule}")

    out_path = OUTPUT_DIR / "jira_workflow_report.docx"
    doc.save(str(out_path))
    print(f"[OK] {out_path}")


# ============================================================
# 3. COST BASELINE — DOCX
# ============================================================
def generate_cost_docx():
    """Генерує cost_baseline_report.docx."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    add_styled_heading(doc, "Cost Baseline", level=0)
    add_styled_heading(doc, "LiLu E-Commerce MVP", level=1)
    add_styled_paragraph(doc, "Автор: Мульков Максим (Dev/QA)", bold=True)
    add_styled_paragraph(doc, "Дата: 06.04.2026 | Версія: 1.0")
    doc.add_paragraph("")

    add_styled_paragraph(doc, 'Наталія попросила: «Прорахуйте мені вартість цього сайту і скажіть, як довго нам чекати до запуску».', italic=True)
    doc.add_paragraph("")

    # 1. Інфраструктура
    add_styled_heading(doc, "1. Вартість інфраструктури (щомісячні витрати)", level=2)
    add_table(doc,
        ["Компонент", "Провайдер", "Вартість (грн/міс)", "Примітка"],
        [
            ["Домен .ua", "Hostiq", "~21 (250/рік)", "Реєстрація на 1 рік"],
            ["SSL-сертифікат", "Let's Encrypt", "0", "Безкоштовний"],
            ["Хостинг Frontend", "Vercel (Hobby)", "0", "Free tier"],
            ["Хостинг Backend", "Railway / Render", "0 – 200", "Free tier достатній"],
            ["База даних", "Render / Supabase", "0 – 150", "Free: 500MB"],
            ["Файлове сховище", "Cloudinary", "0", "Free: 25GB"],
            ["Кешування", "Upstash Redis", "0", "Free: 10K запитів/день"],
            ["Popup-чат", "Tawk.to", "0", "Повністю безкоштовний"],
            ["РАЗОМ", "", "21 – 371", ""],
        ]
    )
    doc.add_paragraph("")

    # 2. Інструменти
    add_styled_heading(doc, "2. Вартість інструментів розробки", level=2)
    add_table(doc,
        ["Інструмент", "Призначення", "Вартість"],
        [
            ["Jira", "Трекінг задач (Free до 10 осіб)", "0"],
            ["GitHub", "Репозиторій коду", "0"],
            ["Figma", "Прототипування", "0"],
            ["VS Code", "Редактор коду", "0"],
            ["Node.js, Next.js, Prisma", "Стек розробки (open source)", "0"],
            ["РАЗОМ", "", "0"],
        ]
    )
    doc.add_paragraph("")

    # 3. Транзакційні
    add_styled_heading(doc, "3. Вартість зовнішніх сервісів", level=2)
    add_table(doc,
        ["Сервіс", "Модель оплати", "Вартість"],
        [
            ["LiqPay (ПриватБанк)", "Комісія з транзакції", "2.75%"],
            ["Нова Пошта API", "Безкоштовний API", "0"],
            ["Кур'єрська доставка", "Залежить від регіону", "Визначає фабрика"],
        ]
    )
    doc.add_paragraph("")

    # 4. Людські ресурси
    add_styled_heading(doc, "4. Вартість людських ресурсів", level=2)
    add_table(doc,
        ["Роль", "Учасник", "Год/тиж", "Тижнів", "Разом годин"],
        [
            ["PM", "Мурадян Руслан", "8", "8", "64"],
            ["BA", "Покась Ілля", "8", "8", "64"],
            ["Dev+QA+SA", "Мульков Максим", "12", "8", "96"],
            ["РАЗОМ", "", "", "", "224"],
        ]
    )
    add_styled_paragraph(doc, "У межах навчального проєкту людські ресурси не тарифікуються. При ставці ~400 грн/год бюджет склав би ≈ 89 600 грн.", italic=True)
    doc.add_paragraph("")

    # 5. Зведення
    add_styled_heading(doc, "5. Загальна вартість запуску MVP", level=2)
    add_table(doc,
        ["Категорія", "Одноразово", "Щомісяця"],
        [
            ["Домен (1 рік)", "~250 грн", "—"],
            ["Хостинг (free tiers)", "0", "0 – 371 грн"],
            ["Інструменти", "0", "0"],
            ["Зовнішні сервіси", "0", "2.75% від продажів"],
            ["Людські ресурси", "0 (навчальний)", "—"],
            ["РАЗОМ", "~250 грн", "0 – 371 грн"],
        ]
    )
    doc.add_paragraph("")
    add_styled_paragraph(doc, "Висновок: Запуск MVP обійдеться від ~250 грн (домен) при використанні безкоштовних планів хостингу. Час до запуску: 2 місяці.", bold=True)

    out_path = OUTPUT_DIR / "cost_baseline_report.docx"
    doc.save(str(out_path))
    print(f"[OK] {out_path}")


# ============================================================
# ЗАПУСК
# ============================================================
if __name__ == "__main__":
    print("[START] Generating documents...\n")
    generate_architecture_md()
    generate_architecture_docx()
    generate_jira_md()
    generate_jira_docx()
    generate_cost_docx()
    print(f"\n[DONE] All files saved to: {OUTPUT_DIR}")
